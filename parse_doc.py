from docx import Document as doc_builder
from docx.document import Document
from typing import List, Tuple
from pprint import pprint
from pathlib import Path
import sys
import re
import os

from jinja2 import (Environment,
                    FileSystemLoader,
                    Template)

def parse_text_number(docx: Document) -> int:
    for p in docx.paragraphs:
        try:
            num = int(p.text)
            return num
        except ValueError:
            continue

def parse_about_paragraphs(docx: Document) -> List[str]:
    about_section_start = False
    paragraphs = []
    for p in docx.paragraphs:
        if about_section_start:
            if re.sub(r'\s', '', p.text.lower()) == 'перевод':
                break
            if p.text:
                paragraphs.append(p.text.strip())
        elif re.sub(r'\s', '', p.text.lower()) == 'отексте':
            about_section_start = True
    return paragraphs

def parse_verses(docx: Document) -> List[Tuple[str, str]]:
    verses = []
    verse_table = docx.tables[0]
    for row in verse_table.rows:
        cells = [cell.text.split('\n') for cell in row.cells]
        if len(cells) > 2: raise ValueError(f"Row had > 2 cells: {cells}")
        lat = [x for x in cells[0] if x]
        ru = [x for x in cells[1] if x]
        if len(lat) != len(ru):
            raise ValueError(f"Verse has different amount of lines: {lat} vs {ru}")
        verses.append((lat, ru))
    return verses

def parse_comments(docx: Document) -> List[Tuple[str | None, str, List[str]]]:
    comments = []
    comment_table = docx.tables[1]
    for row in comment_table.rows:
        ref, text = [cell.text for cell in row.cells]

        ref = ref.split(maxsplit=1)
        num, citate = ref[0], None
        if len(ref) > 1:
            citate = ref[1]

        comments.append((
            citate,
            {'num': num, 'verse': num.split('.')[0]},
            [x.split('\n') for x in text.split('\n\n')]
        ))
    return comments

def parse_bib(docx: Document) -> List[str]:
    bib_section_start = False
    bibs = []
    for p in docx.paragraphs:
        if bib_section_start:
            if 'вступительнойстатьи,переводаикомментария' in re.sub(r'\s', '', p.text.lower()):
                break
            if p.text:
                bibs.append(p.text.strip())
        elif re.sub(r'\s', '', p.text.lower()) == 'библиография':
            bib_section_start = True
    return bibs

def parse_footer(docx: Document) -> Tuple[str, str]:
    names = year = None
    for p in docx.paragraphs:
        if names and year:
            break
        if 'вступительнойстатьи,переводаикомментария' in re.sub(r'\s', '', p.text.lower()):
            names = p.text
        elif 'исследованиеподготовленоприподдержке' in re.sub(r'\s', '', p.text.lower()):
            year = p.text

    return names, year

def parse_docx(file: Path) -> dict:
    data = {}
    docx = doc_builder(file)

    data['text_num'] = parse_text_number(docx)
    data['about'] = parse_about_paragraphs(docx)
    data['verses'] = parse_verses(docx)
    data['commentary'] = parse_comments(docx)
    data['bib'] = parse_bib(docx)
    data['names'], data['year'] = parse_footer(docx)

    return data

def clean_html(html: str) -> str:
    return re.sub(r'\n[ \t]*\n', '\n', html)

def form_html(data: dict, out_file: Path, template_file: Path) -> None:
    template_parent = template_file.parent
    env = Environment(loader=FileSystemLoader(template_parent))
    template = env.get_template(template_file.name)

    with open(out_file, 'w', encoding='utf-8') as f:
        html = template.render(data=data)
        f.write(clean_html(html))

def main():
    if len(sys.argv) < 2:
        print(f"Usage:\npython3 {sys.argv[0]} <input_file> <output_file>")
        input_file, output_file = Path("/home/kesha/Downloads/Telegram Desktop/217 -.docx"), Path("texts/217.html")
        #exit(1)
    else:
        input_file, output_file = Path(sys.argv[1]), Path(sys.argv[2])
    template_file = Path(__file__).parent.joinpath('templates/text.html')
    if not input_file.is_file():
        raise FileNotFoundError(f"File {input_file} does not exist")

    data = parse_docx(input_file)
    form_html(data, output_file, template_file)
    
if __name__ == '__main__':
    main()
